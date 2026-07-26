import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

def get_explanation(filename, content):
    lines = content.split('\n')
    
    extends_class = ""
    class_name = ""
    has_ready = False
    has_process = False
    signals = []
    
    for line in lines:
        line_strip = line.strip()
        if line_strip.startswith('extends '):
            extends_class = line_strip.split(' ')[1]
        elif line_strip.startswith('class_name '):
            class_name = line_strip.split(' ')[1]
        elif line_strip.startswith('func _ready()'):
            has_ready = True
        elif line_strip.startswith('func _process('):
            has_process = True
        elif line_strip.startswith('signal '):
            sig_name = line_strip.split(' ')[1].split('(')[0]
            signals.append(sig_name)
    
    desc_parts = [f"Script {filename} berfungsi sebagai pengontrol logika dan interaksi untuk bagian {filename.replace('.gd', '')}."]
    
    if extends_class:
        class_str = f" yang didefinisikan sebagai class `{class_name}`" if class_name else ""
        desc_parts.append(f"Script ini mewarisi properti dan metode dari class bawaan `{extends_class}`{class_str}.")
        
    if has_ready or has_process:
        func_parts = []
        if has_ready:
            func_parts.append("`_ready()` untuk melakukan inisialisasi awal saat node pertama kali masuk ke dalam scene tree")
        if has_process:
            func_parts.append("`_process()` untuk menangani pembaruan logika yang berjalan setiap frame secara real-time")
        desc_parts.append(f"Dalam implementasi kodenya, script ini memanfaatkan fungsi {' serta fungsi '.join(func_parts)}.")
        
    if signals:
        desc_parts.append(f"Selain itu, script ini juga dilengkapi dengan custom signal ({', '.join(['`'+s+'`' for s in signals])}) yang bertugas memancarkan notifikasi atau memicu event pada node lain ketika kondisi tertentu terpenuhi.")
        
    desc_parts.append("Secara keseluruhan, kumpulan baris kode ini memastikan fungsionalitas dan mekanik terkait dapat terintegrasi serta merespons dengan baik di dalam ekosistem game.")
    
    full_paragraph = " ".join(desc_parts)
    return full_paragraph

def main():
    dir_path = r'c:\Users\ASUS\OneDrive\Documents\GitHub\vn-jika-kucing-lenyap-dari-dunia-(4.6)\scripts'
    output_path = r'c:\Users\ASUS\OneDrive\Documents\GitHub\vn-jika-kucing-lenyap-dari-dunia-(4.6)\Penjelasan_Scripts_v2.docx'
    
    document = Document()
    
    # Set default font to Calibri
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    document.add_heading('Penjelasan File Script (res://scripts/)', 0)
    
    files = [f for f in os.listdir(dir_path) if f.endswith('.gd')]
    files.sort()
    
    for idx, f in enumerate(files, 1):
        filepath = os.path.join(dir_path, f)
        with open(filepath, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # 1. Nama File (Bold, Hitam, Calibri)
        p_title = document.add_paragraph()
        run_title = p_title.add_run(f"{idx}. {f}")
        run_title.bold = True
        run_title.font.name = 'Calibri'
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        run_title.font.size = Pt(14)
        
        # Prepare content with line numbers
        lines = content.split('\n')
        numbered_lines = []
        for i, line in enumerate(lines, 1):
            numbered_lines.append(f"{i:3d} | {line}")
        numbered_content = '\n'.join(numbered_lines)
        
        # 2. [Copy full script] dalam tabel (kotak) dengan Courier New
        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        
        p_cell = cell.paragraphs[0]
        run_cell = p_cell.add_run(numbered_content)
        run_cell.font.name = 'Courier New'
        run_cell.font.size = Pt(9)
        
        document.add_paragraph() # Add spacing
        
        # 3. Penjelasan kegunaan script dan baris penting
        desc = get_explanation(f, content)
        
        p_desc = document.add_paragraph()
        run_desc_title = p_desc.add_run('Penjelasan Kegunaan Script & Baris Penting:\n')
        run_desc_title.bold = True
        run_desc_title.font.name = 'Calibri'
        
        run_desc = p_desc.add_run(desc)
        run_desc.font.name = 'Calibri'
        
        document.add_paragraph() # Add spacing before next
        
        document.add_paragraph() # Add spacing before next
        
    document.save(output_path)
    print(f"File saved to {output_path}")

if __name__ == '__main__':
    main()
