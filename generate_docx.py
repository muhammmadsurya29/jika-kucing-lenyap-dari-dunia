import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

def main():
    dir_path = r'c:\Users\ASUS\OneDrive\Documents\GitHub\vn-jika-kucing-lenyap-dari-dunia-(4.6)\dialogues'
    output_path = r'c:\Users\ASUS\OneDrive\Documents\GitHub\vn-jika-kucing-lenyap-dari-dunia-(4.6)\Penjelasan_Dialog_v2.docx'
    
    document = Document()
    
    # Set default font to Calibri
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    document.add_heading('Penjelasan File Dialog (res://dialogues/)', 0)
    
    files = [f for f in os.listdir(dir_path) if f.endswith('.dtl')]
    files.sort()
    
    for idx, f in enumerate(files, 1):
        filepath = os.path.join(dir_path, f)
        with open(filepath, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Parse elements
        bg_matches = re.findall(r'\[background[^\]]*\]', content)
        join_matches = re.findall(r'join\s+[^\s]+', content)
        audio_matches = re.findall(r'\[music[^\]]*\]|\[sound[^\]]*\]', content)
        signal_matches = re.findall(r'\[signal[^\]]*\]', content)
        
        bgs = set(bg_matches + join_matches)
        audios = set(audio_matches)
        signals = set(signal_matches)
        
        bg_str = ', '.join(bgs) if bgs else 'Tidak ada'
        audio_str = ', '.join(audios) if audios else 'Tidak ada'
        signal_str = ', '.join(signals) if signals else 'Tidak ada'
        
        # 1. Nama File (Bold, Hitam, Calibri)
        p_title = document.add_paragraph()
        run_title = p_title.add_run(f"{idx}. {f}")
        run_title.bold = True
        run_title.font.name = 'Calibri'
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        run_title.font.size = Pt(14) # Make it slightly larger like a heading
        
        # [Copy full script] dalam tabel (kotak)
        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        
        # Add content into cell
        p_cell = cell.paragraphs[0]
        run_cell = p_cell.add_run(content)
        run_cell.font.name = 'Courier New' # Script text
        run_cell.font.size = Pt(9)
        
        document.add_paragraph() # Add spacing
        
        # Helper function to add bullets with Calibri
        def add_bullet_item(bold_text, normal_text):
            p = document.add_paragraph(style='List Bullet')
            run_bold = p.add_run(bold_text)
            run_bold.bold = True
            run_bold.font.name = 'Calibri'
            run_normal = p.add_run(normal_text)
            run_normal.font.name = 'Calibri'
            
        # - Penjelasan
        add_bullet_item('Penjelasan kegunaan script dialog: ', f'File ini berisi logika dan percakapan untuk bagian/scene {f.replace(".dtl", "")}.')
        
        # - Background, signal, dll
        add_bullet_item('Background, Karakter (Join): ', bg_str)
        add_bullet_item('Audio (Music/Sound): ', audio_str)
        add_bullet_item('Signal (Event): ', signal_str)
        
        document.add_paragraph() # Add spacing before next
        
    document.save(output_path)
    print(f"File saved to {output_path}")

if __name__ == '__main__':
    main()
